import docx
doc = docx.Document()
doc.add_paragraph("Mr. John Smith and Ms. Jane Doe")
doc.add_paragraph("invite you to join them as they exchange vows on")
doc.add_paragraph("Saturday, the tenth of June, two thousand twenty-two")
doc.add_paragraph("at five o'clock in the evening")
doc.add_paragraph("St. Mark's Church")
doc.add_paragraph("123 Main Street")
doc.add_paragraph("Anytown, USA")

doc.save("marrage.docx")
